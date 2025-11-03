# app3.py  -- Advanced RAG Q&A (upgraded)
# Fixed ChromaDB permissions issue

import os
import tempfile
import streamlit as st
import time
import shutil
import atexit
import uuid
import json
from datetime import datetime
from dotenv import load_dotenv

# LangChain / embedding / vectorstore imports
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser

# Optional imports
try:
    from sklearn.cluster import KMeans
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Load environment variables
load_dotenv()

# -----------------------------
# Streamlit basic config / UI
# -----------------------------
st.set_page_config(page_title="Advanced RAG Q&A", page_icon="📚", layout="wide")
st.title("🚀 Advanced RAG Q&A — Upgraded")

# Debug toggle
DEBUG = st.sidebar.checkbox("Debug mode (console logs)", False)

def debug_log(msg):
    if DEBUG:
        print("DEBUG:", msg)

# -----------------------------
# Session & storage helpers
# -----------------------------
# Use temp directory for better compatibility in deployments
DEFAULT_BASE_DIR = tempfile.gettempdir()
os.makedirs(DEFAULT_BASE_DIR, exist_ok=True)

def session_dir(session_id):
    # Use a more portable path structure
    safe_session_id = "".join(c for c in session_id if c.isalnum() or c in ('-', '_')).rstrip()
    d = os.path.join(DEFAULT_BASE_DIR, "rag_sessions", f"session_{safe_session_id}")
    os.makedirs(d, exist_ok=True)
    return d

def chat_history_path(session_id):
    return os.path.join(session_dir(session_id), "chat_history.json")

def save_chat_history(session_id, messages):
    try:
        with open(chat_history_path(session_id), "w", encoding="utf-8") as f:
            json.dump({"messages": messages, "updated_at": datetime.utcnow().isoformat()}, f, ensure_ascii=False, indent=2)
        debug_log(f"Saved chat history for {session_id}")
    except Exception as e:
        debug_log(f"Error saving chat history: {e}")

def load_chat_history(session_id):
    p = chat_history_path(session_id)
    if os.path.exists(p):
        try:
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("messages", [])
        except Exception as e:
            debug_log(f"Error loading chat history: {e}")
    return []

# -----------------------------
# Sidebar configuration
# -----------------------------
with st.sidebar:
    st.header("⚙️ Configuration")
    session_id = st.text_input("Session ID", value=os.getenv("DEFAULT_SESSION_ID", "default_session"))
    api_key_input = st.text_input("Groq API Key", type="password")
    model_choice = st.selectbox("Model", ["llama-3.3-70b-versatile", "llama-3.3-70b", "openai/gpt-oss-120b"], index=0)
    top_k = st.slider("Top K Chunks", 1, 12, 4)
    similarity_threshold = st.slider("Similarity Threshold (distance)", 0.0, 1.0, 0.7)
    debug_checkbox = st.checkbox("Show more debug info", False)
    
    # Enhanced cleanup button
    if st.button("🧹 Cleanup Vector Store (this session)"):
        idx_dir = os.path.join(session_dir(session_id), "chroma_index")
        if os.path.exists(idx_dir):
            shutil.rmtree(idx_dir, ignore_errors=True)
            st.success("Cleaned session vector store.")
            # Clear all related cache
            keys_to_remove = [k for k in st.session_state.keys() if 'vectorstore' in k]
            for key in keys_to_remove:
                del st.session_state[key]
            st.cache_resource.clear()
            st.rerun()
        else:
            st.info("No vectorstore found for this session.")

# link debug flags
if debug_checkbox:
    DEBUG = True

# Use API key provided or from .env
api_key = api_key_input or os.getenv("GROQ_API_KEY")
if not api_key:
    st.warning("⚠️ Please enter your Groq API key in the sidebar or in a .env file.")
    st.stop()

# Initialize LLM
try:
    llm = ChatGroq(
        groq_api_key=api_key,
        model_name=model_choice,
        temperature=0.1
    )
except Exception as e:
    st.error(f"Failed to initialize Groq model: {e}")
    st.stop()

# -----------------------------
# File upload (PDF, txt, docx)
# -----------------------------
st.header("📤 Upload documents (PDF / TXT / DOCX)")
uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx"], accept_multiple_files=True)

# Track uploaded files to detect changes
if 'previous_uploaded_files' not in st.session_state:
    st.session_state.previous_uploaded_files = []

# Check if files have changed
current_files = [f.name for f in uploaded_files] if uploaded_files else []
previous_files = st.session_state.previous_uploaded_files

files_changed = current_files != previous_files

if files_changed:
    st.session_state.previous_uploaded_files = current_files
    # Clear vectorstore cache when files change
    keys_to_remove = [k for k in st.session_state.keys() if 'vectorstore' in k]
    for key in keys_to_remove:
        del st.session_state[key]

if not uploaded_files:
    st.info("Upload one or more files to begin. You can also drag multiple files.")
    # Load previous chat history and show quick actions if available
    previous_msgs = load_chat_history(session_id)
    if previous_msgs:
        st.success(f"Loaded saved chat for session '{session_id}' ({len(previous_msgs)} messages).")
        if st.button("🔁 Load chat into UI"):
            st.session_state.messages = previous_msgs
    st.stop()

# -----------------------------
# Temporary save uploaded files and load docs
# -----------------------------
# Document loader: prefer PyMuPDFLoader; fallback to PyPDFLoader
try:
    from langchain_community.document_loaders import PyMuPDFLoader as PDFLoader
except Exception:
    from langchain_community.document_loaders import PyPDFLoader as PDFLoader

from langchain_core.documents import Document

all_docs = []
tmp_paths = []
with st.spinner("Processing uploaded files..."):
    for f in uploaded_files:
        name = f.name
        suffix = os.path.splitext(name)[1].lower()
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(f.getvalue())
            tmp_paths.append(tmp.name)
            tmp_name = tmp.name

        try:
            if suffix == ".pdf":
                loader = PDFLoader(tmp_name)
                docs = loader.load()
            elif suffix == ".txt":
                text = open(tmp_name, "r", encoding="utf-8", errors="ignore").read()
                docs = [Document(page_content=text, metadata={"source_file": name})]
            elif suffix == ".docx" and DOCX_AVAILABLE:
                doc = docx.Document(tmp_name)
                full = "\n".join([p.text for p in doc.paragraphs])
                docs = [Document(page_content=full, metadata={"source_file": name})]
            else:
                # fallback: plain read
                text = open(tmp_name, "r", encoding="utf-8", errors="ignore").read()
                docs = [Document(page_content=text, metadata={"source_file": name})]

            # ensure metadata tags
            for d in docs:
                d.metadata["source_file"] = name
                d.metadata["upload_time"] = datetime.utcnow().isoformat()
            all_docs.extend(docs)
        except Exception as e:
            st.error(f"Failed to load {name}: {e}")

# cleanup temp files
for p in tmp_paths:
    try:
        os.unlink(p)
    except Exception:
        pass

if not all_docs:
    st.error("No readable text found in uploaded documents. Try different files.")
    st.stop()

st.success(f"✅ Loaded {len(all_docs)} pages/chunks from {len(uploaded_files)} file(s)")

# -----------------------------
# Split into chunks
# -----------------------------
splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=120)
splits = splitter.split_documents(all_docs)
st.info(f"📄 Created {len(splits)} text chunks.")

# -----------------------------
# Embeddings (cache)
# -----------------------------
@st.cache_resource(show_spinner=False)
def get_embeddings():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

embeddings = get_embeddings()

# -----------------------------
# Vectorstore - FIXED VERSION with better error handling
# -----------------------------
INDEX_DIR = os.path.join(session_dir(session_id), "chroma_index")

def init_vectorstore_robust(_splits, _embeddings, persist_dir):
    """
    Robust vectorstore initialization with multiple fallback strategies
    """
    strategies = [
        # Strategy 1: Try with persistence first
        lambda: Chroma.from_documents(_splits, _embeddings, persist_directory=persist_dir),
        # Strategy 2: Try without persistence if above fails
        lambda: Chroma.from_documents(_splits, _embeddings),
        # Strategy 3: Try in-memory only
        lambda: Chroma.from_documents(_splits, _embeddings, persist_directory=None)
    ]
    
    last_error = None
    for i, strategy in enumerate(strategies):
        try:
            debug_log(f"Trying vectorstore strategy {i+1}")
            vs = strategy()
            debug_log(f"Success with strategy {i+1}")
            return vs
        except Exception as e:
            last_error = e
            debug_log(f"Strategy {i+1} failed: {e}")
            # Clean up before next attempt
            if os.path.exists(persist_dir):
                try:
                    shutil.rmtree(persist_dir, ignore_errors=True)
                except:
                    pass
            continue
    
    # If all strategies fail, raise the last error
    raise RuntimeError(f"All vectorstore initialization strategies failed. Last error: {last_error}")

# Initialize vectorstore
vectorstore_key = f"vectorstore_{session_id}_{hash(tuple(current_files)) if current_files else 'no_files'}"

if vectorstore_key not in st.session_state or files_changed:
    with st.spinner("Initializing vector store..."):
        try:
            # Ensure directory permissions
            os.makedirs(INDEX_DIR, exist_ok=True)
            
            # Test write permissions
            test_file = os.path.join(INDEX_DIR, "write_test.txt")
            try:
                with open(test_file, "w") as f:
                    f.write("test")
                os.remove(test_file)
            except Exception as e:
                debug_log(f"Write test failed: {e}")
                st.warning("⚠️ Limited write permissions detected. Using in-memory vector store.")
            
            st.session_state[vectorstore_key] = init_vectorstore_robust(splits, embeddings, INDEX_DIR)
            st.session_state.vectorstore_initialized = True
            debug_log("Vectorstore initialized successfully")
            
        except Exception as e:
            st.error(f"Vectorstore initialization error: {e}")
            st.info("💡 **Troubleshooting tips:**")
            st.info("1. Try using the 'Cleanup Vector Store' button in the sidebar")
            st.info("2. Try a different Session ID")
            st.info("3. The app will use in-memory storage for this session")
            st.stop()
else:
    debug_log("Using cached vectorstore")

vectorstore = st.session_state[vectorstore_key]

# -----------------------------
# Retriever with threshold filtering
# -----------------------------
def create_retriever(vs, k=5, distance_threshold=0.7):
    def retrieve(question):
        try:
            docs_with_scores = vs.similarity_search_with_score(question, k=k*3)
            filtered = [d for d, dist in docs_with_scores if dist <= distance_threshold]
            if not filtered:
                filtered = [d for d, _ in docs_with_scores[:k]]
            return filtered[:k]
        except Exception as e:
            debug_log(f"Retriever internal error: {e}")
            return vs.similarity_search(question, k=k)
    return retrieve

retriever = create_retriever(vectorstore, k=top_k, distance_threshold=similarity_threshold)

# -----------------------------
# RAG chain definition
# -----------------------------
def rag_chain(question, chat_history_messages):
    start_time = time.time()
    docs = retriever(question)
    if not docs:
        return "⚠️ No relevant text found in your documents.", [], 0.0

    context = "\n\n".join([d.page_content for d in docs])
    qa_prompt = ChatPromptTemplate.from_messages([
        ("system",
         "You are a helpful assistant. Answer based **only** on the provided context.\n"
         "If the answer is not in the context, say: 'The answer is not available in the provided documents.'\n\n"
         "Context:\n{context}\n"),
        ("human", "Question: {question}")
    ])

    chain = qa_prompt | llm | StrOutputParser()
    try:
        response = chain.invoke({"context": context, "question": question})
    except Exception as e:
        response = f"Error generating response: {e}"
    duration = round(time.time() - start_time, 2)
    return response, docs, duration

# -----------------------------
# Insights & clustering
# -----------------------------
def generate_document_summary(llm, docs, max_chars=4000):
    preview = "\n\n".join([d.page_content[:1000] for d in docs[:10]])
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a helpful assistant summarizer."),
        ("human", f"Summarize the following documents into a concise set of bullet points (max 6 bullets). Text:\n\n{preview}")
    ])
    chain = prompt | llm | StrOutputParser()
    try:
        return chain.invoke({})
    except Exception as e:
        return f"Unable to generate summary: {e}"

def cluster_chunks(embeddings, chunks, n_clusters=4):
    if not SKLEARN_AVAILABLE:
        return None
    try:
        embs = [embeddings.embed_query(c.page_content) for c in chunks]
        import numpy as np
        em = np.array(embs)
        kmeans = KMeans(n_clusters=min(n_clusters, len(chunks)), random_state=0).fit(em)
        clusters = {}
        for idx, label in enumerate(kmeans.labels_):
            clusters.setdefault(int(label), []).append(chunks[idx])
        return clusters
    except Exception as e:
        debug_log(f"Clustering failed: {e}")
        return None

# -----------------------------
# Chat UI & interaction
# -----------------------------
if "messages" not in st.session_state:
    previous = load_chat_history(session_id)
    if previous:
        st.session_state.messages = previous
    else:
        st.session_state.messages = []

if "chathistory" not in st.session_state:
    st.session_state.chathistory = {}

def get_chat_history_obj(sid):
    if sid not in st.session_state.chathistory:
        st.session_state.chathistory[sid] = ChatMessageHistory()
    return st.session_state.chathistory[sid]

# Display existing messages
st.header("💬 Chat with your documents")
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        st.markdown(m["content"])

# Show insights block
with st.expander("🔎 Document Insights"):
    st.write(f"**Files uploaded:** {', '.join({d.metadata.get('source_file','?') for d in all_docs})}")
    st.write(f"**Total raw pages/chunks:** {len(all_docs)}")
    st.write(f"**Chunks created:** {len(splits)}")
    
    if files_changed:
        st.success("🔄 Documents updated - vector store refreshed!")
    
    if st.button("🧾 Generate quick summary (using model)"):
        with st.spinner("Generating summary..."):
            summary = generate_document_summary(llm, all_docs)
            st.markdown(summary)

    if SKLEARN_AVAILABLE:
        if st.button("📂 Try topic clustering of chunks"):
            with st.spinner("Clustering chunks..."):
                clusters = cluster_chunks(embeddings, splits, n_clusters=4)
                if clusters:
                    for k, items in clusters.items():
                        st.markdown(f"**Cluster {k} — {len(items)} chunks**")
                        st.write(items[0].page_content[:300] + "...")
                else:
                    st.info("Clustering did not produce results.")
    else:
        st.info("Topic clustering requires scikit-learn (`pip install scikit-learn`).")

# Chat input
user_q = st.chat_input("Ask something about your uploaded documents...")

if user_q:
    history_obj = get_chat_history_obj(session_id)
    with st.chat_message("user"):
        st.markdown(user_q)
    st.session_state.messages.append({"role": "user", "content": user_q})
    history_obj.add_user_message(user_q)

    with st.chat_message("assistant"):
        with st.spinner("Analyzing documents and generating answer..."):
            try:
                response, docs, duration = rag_chain(user_q, history_obj)
                st.markdown(response)
                st.caption(f"⏱️ Answer generated in {duration} seconds")

                if docs:
                    with st.expander("📂 Sources Used (click to expand)"):
                        for d in docs:
                            src = d.metadata.get("source_file", "Unknown File")
                            st.markdown(f"**📄 {src}** — Preview:")
                            snippet = d.page_content[:800]
                            try:
                                import re
                                pattern = re.compile(re.escape(user_q[:60]), re.IGNORECASE)
                                snippet = pattern.sub(lambda m: f"**{m.group(0)}**", snippet)
                            except Exception:
                                pass
                            st.write(snippet + "...")
                            st.divider()

                st.session_state.messages.append({"role": "assistant", "content": response})
                history_obj.add_ai_message(response)
                save_chat_history(session_id, st.session_state.messages)

            except Exception as e:
                err = f"Error while generating response: {e}"
                st.error(err)
                st.session_state.messages.append({"role": "assistant", "content": err})
                save_chat_history(session_id, st.session_state.messages)

# -----------------------------
# Chat export / download
# -----------------------------
st.header("📥 Export / Session Tools")
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("📥 Download Chat as TXT"):
        chat_text = "\n".join([f"{m['role'].upper()}: {m['content']}" for m in st.session_state.messages])
        st.download_button("Download TXT", data=chat_text, file_name=f"chat_{session_id}.txt", mime="text/plain")

with col2:
    if st.button("💾 Export Chat as JSON"):
        chat_json = json.dumps(st.session_state.messages, ensure_ascii=False, indent=2)
        st.download_button("Download JSON", data=chat_json, file_name=f"chat_{session_id}.json", mime="application/json")

with col3:
    if st.button("🗑️ Clear current chat (session)"):
        st.session_state.messages = []
        p = chat_history_path(session_id)
        if os.path.exists(p):
            os.remove(p)
        st.rerun()

# -----------------------------
# Admin & diagnostics
# -----------------------------
with st.expander("📊 Admin / Diagnostics"):
    st.write(f"Vectorstore directory: {INDEX_DIR}")
    try:
        idx_exists = os.path.exists(INDEX_DIR)
        st.write(f"Vectorstore exists: {idx_exists}")
        if idx_exists:
            doc_count = vectorstore._collection.count()
            st.write(f"Documents in vectorstore: {doc_count}")
    except Exception as e:
        st.write(f"Error checking vectorstore: {e}")
    
    st.write(f"Current uploaded files: {current_files}")
    st.write(f"Files changed since last run: {files_changed}")
    st.write(f"Session ID: {session_id}")

# -----------------------------
# Atexit cleanup
# -----------------------------
def cleanup_vectorstore_on_exit():
    debug_log("Exiting app. Manual cleanup available via sidebar.")

atexit.register(cleanup_vectorstore_on_exit)
